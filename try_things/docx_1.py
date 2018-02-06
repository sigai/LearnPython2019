from docx import Document
from docx.shared import Inches

# 新建document对象
document = Document()

# 添加段落对象
paragraph = document.add_paragraph("hello world".title())

# 插入段落
paragraph.insert_paragraph_before("Python")

# 添加heading
document.add_heading("this is default heading")

# 添加换页
# document.add_page_break()

# 添加表格
table = document.add_table(rows=2, cols=3)
table.add_row() #添加行

for row in table.rows:      # 遍历表格
    for cell in row.cells:
        cell.text = "fuck"
cell = table.cell(0,0)  #单元格

# 添加图像并调整大小
document.add_picture("test.gif", width=Inches(1.0))

# 样式
paragraph = document.add_paragraph("Did i looking better?")
paragraph.style = "ListBullet"

# run
paragraph = document.add_paragraph("this is before run test ")
run = paragraph.add_run("this is test run")
run.bold = True
run.style = "Emphasis"  # 样式
paragraph.add_run(" run ends hear.")

# 保存文档
document.save("test.docx")
